"""
Document processing utilities for handling different file types
"""
import io
import logging
from typing import Optional, Union
from pathlib import Path
import PyPDF2
import docx
from app.core.exceptions.exceptions import TextExtractionError, FileProcessingError
from app.config.settings import settings

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles document processing for different file types"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._extract_pdf_text,
            '.docx': self._extract_docx_text,
            '.txt': self._extract_txt_text
        }
    
    def extract_text(
        self,
        file_bytes: bytes,
        filename: str,
        file_extension: Optional[str] = None
    ) -> str:
        """
        Extract text from document bytes
        
        Args:
            file_bytes: Raw file bytes
            filename: Original filename
            file_extension: File extension (optional, will be extracted from filename if not provided)
            
        Returns:
            Extracted text content
            
        Raises:
            TextExtractionError: If text extraction fails
            FileProcessingError: If file format is not supported
        """
        try:
            if not file_extension:
                file_extension = Path(filename).suffix.lower()
            
            if file_extension not in self.supported_formats:
                raise FileProcessingError(
                    f"Unsupported file format: {file_extension}",
                    error_code="UNSUPPORTED_FORMAT",
                    details={"supported_formats": list(self.supported_formats.keys())}
                )
            
            logger.info(f"Processing {filename} with format {file_extension}")
            extractor = self.supported_formats[file_extension]
            text = extractor(file_bytes)
            
            if not text.strip():
                logger.warning(f"Extracted text is empty for {filename}")
                return ""
            
            logger.info(f"Successfully extracted {len(text)} characters from {filename}")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Failed to extract text from {filename}: {str(e)}")
            if isinstance(e, (TextExtractionError, FileProcessingError)):
                raise
            raise TextExtractionError(
                f"Failed to extract text from {filename}",
                error_code="EXTRACTION_FAILED",
                details={"original_error": str(e)}
            )
    
    def _extract_pdf_text(self, file_bytes: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text_parts = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(page_text)
                    else:
                        logger.warning(f"Page {page_num + 1} appears to be empty or unreadable")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {str(e)}")
                    continue
            
            if not text_parts:
                raise TextExtractionError("No readable text found in PDF")
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"PDF text extraction failed: {str(e)}")
            raise TextExtractionError(f"PDF text extraction failed: {str(e)}")
    
    def _extract_docx_text(self, file_bytes: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            if not text_parts:
                raise TextExtractionError("No readable text found in DOCX")
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"DOCX text extraction failed: {str(e)}")
            raise TextExtractionError(f"DOCX text extraction failed: {str(e)}")
    
    def _extract_txt_text(self, file_bytes: bytes) -> str:
        """Extract text from plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = file_bytes.decode(encoding)
                    return text
                except UnicodeDecodeError:
                    continue
            
            raise TextExtractionError("Failed to decode text file with any supported encoding")
            
        except Exception as e:
            logger.error(f"Text file extraction failed: {str(e)}")
            raise TextExtractionError(f"Text file extraction failed: {str(e)}")
    
    def validate_file_size(self, file_bytes: bytes, max_size: Optional[int] = None) -> bool:
        """Validate file size"""
        max_size = max_size or settings.MAX_FILE_SIZE
        file_size = len(file_bytes)
        
        if file_size > max_size:
            logger.warning(f"File size {file_size} exceeds limit {max_size}")
            return False
        
        return True
    
    def get_file_info(self, file_bytes: bytes, filename: str) -> dict:
        """Get basic file information"""
        file_extension = Path(filename).suffix.lower()
        file_size = len(file_bytes)
        
        return {
            "filename": filename,
            "extension": file_extension,
            "size_bytes": file_size,
            "size_mb": round(file_size / (1024 * 1024), 2),
            "is_supported": file_extension in self.supported_formats
        }
